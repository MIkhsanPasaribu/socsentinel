"""
SOCsentinel — DOCX Report Generator.

Generates professional Word documents using python-docx.
"""

from io import BytesIO
from typing import Any

from app.core.logger import get_logger

logger = get_logger(__name__)


def _check_python_docx() -> bool:
    """Check whether python-docx can be imported at runtime."""
    try:
        import docx  # noqa: F401
        return True
    except ImportError:
        return False


DOCX_AVAILABLE = _check_python_docx()


def generate_docx(context: dict[str, Any]) -> BytesIO:
    """Generate a DOCX report from the report context.

    Creates a professional Word document with all investigation details,
    formatted for SOC teams to edit and share.

    Args:
        context: The report context dictionary with all investigation data.

    Returns:
        BytesIO buffer containing the generated DOCX.

    Raises:
        RuntimeError: If python-docx is not installed.
    """
    if not DOCX_AVAILABLE:
        raise RuntimeError(
            "DOCX generation is not available. "
            "Install python-docx: pip install python-docx"
        )

    # Lazy imports — only reached when python-docx is confirmed available
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # SOCsentinel brand colors
    COLOR_DEEP_NAVY = RGBColor(0x0A, 0x16, 0x28)
    COLOR_CYAN = RGBColor(0x00, 0xD4, 0xFF)
    COLOR_GRAY = RGBColor(0x6B, 0x72, 0x80)

    def _severity_color(severity: str) -> RGBColor:
        """Get RGB color for severity level."""
        s = severity.lower()
        if s == "critical":
            return RGBColor(0xDC, 0x26, 0x26)
        elif s == "high":
            return RGBColor(0xEA, 0x58, 0x0C)
        elif s == "medium":
            return RGBColor(0xD9, 0x77, 0x06)
        elif s == "low":
            return RGBColor(0x05, 0x96, 0x69)
        return COLOR_GRAY

    def _add_heading(doc: Document, text: str, level: int = 1) -> Any:
        """Add a styled heading to the document."""
        heading = doc.add_heading(text, level=level)
        for run in heading.runs:
            run.font.color.rgb = COLOR_DEEP_NAVY
            if level == 1:
                run.font.size = Pt(18)
            elif level == 2:
                run.font.size = Pt(14)
            else:
                run.font.size = Pt(12)
            run.font.bold = True
        return heading

    def _add_info_table(doc: Document, items: list[tuple[str, str]]) -> None:
        """Add a two-column info table."""
        table = doc.add_table(rows=len(items), cols=2)
        table.style = "Table Grid"
        for i, (label, value) in enumerate(items):
            row = table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = value
            row.cells[0].width = Inches(1.5)
            for paragraph in row.cells[0].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.size = Pt(10)
                    run.font.color.rgb = COLOR_GRAY
            for paragraph in row.cells[1].paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(10)

    # ── Build document ──────────────────────────────────────────────

    doc = Document()

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)

    # Narrow margins
    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # === COVER PAGE ===
    for _ in range(3):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("SOCsentinel")
    run.font.size = Pt(36)
    run.font.bold = True
    run.font.color.rgb = COLOR_DEEP_NAVY

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run("Security Investigation Report")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = COLOR_CYAN

    for _ in range(2):
        doc.add_paragraph()

    sev_para = doc.add_paragraph()
    sev_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sev_run = sev_para.add_run(context["alert"]["severity"].upper())
    sev_run.font.size = Pt(16)
    sev_run.font.bold = True
    sev_run.font.color.rgb = _severity_color(context["alert"]["severity"])

    for _ in range(2):
        doc.add_paragraph()

    for label, value in [
        ("Investigation ID:", context["investigation_id"]),
        ("Alert:", context["alert"]["rule_name"]),
        ("Generated:", context["generated_at_formatted"]),
        ("Processing Time:", f"{context['total_processing_time_ms']:.0f}ms"),
    ]:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lr = para.add_run(f"{label} ")
        lr.font.size = Pt(11)
        lr.font.color.rgb = COLOR_GRAY
        vr = para.add_run(value)
        vr.font.size = Pt(11)
        vr.font.bold = True

    doc.add_page_break()

    # === EXECUTIVE SUMMARY ===
    _add_heading(doc, "Executive Summary", level=1)

    p = doc.add_paragraph()
    r = p.add_run(context["report"]["title"])
    r.font.bold = True
    r.font.size = Pt(12)
    doc.add_paragraph(context["report"]["executive_summary"])

    _add_heading(doc, "Key Metrics", level=2)
    for label, value in [
        ("AI Confidence", f"{context['overall_confidence'] * 100:.0f}%"),
        ("Classification", context["triage"]["classification"].upper()),
        ("MITRE Techniques", str(len(context["mitre"]["techniques"]))),
        ("IOCs Found", str(len(context["evidence"]["iocs"]))),
    ]:
        bp = doc.add_paragraph(style="List Bullet")
        bp.add_run(f"{label}: ").bold = True
        bp.add_run(value)

    if context["report"]["recommendations"]:
        _add_heading(doc, "Key Recommendations", level=2)
        for rec in context["report"]["recommendations"][:5]:
            doc.add_paragraph(rec, style="List Bullet")

    # === ALERT DETAILS ===
    doc.add_page_break()
    _add_heading(doc, "Alert Details", level=1)
    _add_info_table(doc, [
        ("Alert ID", context["alert"]["alert_id"]),
        ("Rule Name", context["alert"]["rule_name"]),
        ("Source IP", context["alert"]["source_ip"] or "N/A"),
        ("Destination IP", context["alert"]["destination_ip"] or "N/A"),
        ("Username", context["alert"]["username"] or "N/A"),
        ("Hostname", context["alert"]["hostname"] or "N/A"),
        ("Protocol", context["alert"]["protocol"] or "N/A"),
        ("Timestamp", context["alert"]["timestamp"]),
    ])

    if context["alert"]["description"]:
        doc.add_paragraph()
        dp = doc.add_paragraph()
        dp.add_run("Description: ").bold = True
        doc.add_paragraph(context["alert"]["description"])

    # === TRIAGE ANALYSIS ===
    _add_heading(doc, "Triage Analysis (L1)", level=1)
    _add_info_table(doc, [
        ("Classification", context["triage"]["classification"].upper()),
        ("Severity", context["triage"]["severity"].upper()),
        ("AI Confidence", f"{context['triage']['confidence'] * 100:.0f}%"),
        ("False Positive Probability", f"{context['triage']['false_positive_probability'] * 100:.0f}%"),
        ("False Positive", "Yes" if context["triage"]["is_false_positive"] else "No"),
    ])

    if context["triage"]["reasoning"]:
        doc.add_paragraph()
        rp = doc.add_paragraph()
        rp.add_run("Reasoning: ").bold = True
        doc.add_paragraph(context["triage"]["reasoning"])

    if context["triage"]["evidence_chain"]:
        _add_heading(doc, "Evidence Chain", level=2)
        for step in context["triage"]["evidence_chain"]:
            bp = doc.add_paragraph(style="List Bullet")
            bp.add_run(f"{step['step']}: ").bold = True
            bp.add_run(f"{step['observation']} -> {step['conclusion']}")

    # === INDICATORS OF COMPROMISE ===
    if context["evidence"]["iocs"]:
        doc.add_page_break()
        _add_heading(doc, "Indicators of Compromise", level=1)

        ioc_table = doc.add_table(rows=1, cols=4)
        ioc_table.style = "Table Grid"
        ioc_table.autofit = True
        for i, h in enumerate(["Type", "Value", "Reputation", "Source"]):
            ioc_table.rows[0].cells[i].text = h
            for paragraph in ioc_table.rows[0].cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = COLOR_DEEP_NAVY
                    run.font.size = Pt(10)

        for ioc in context["evidence"]["iocs"]:
            cells = ioc_table.add_row().cells
            cells[0].text = ioc.get("type", "").upper()
            cells[1].text = ioc.get("value", "")
            cells[2].text = ioc.get("reputation", "unknown").upper()
            cells[3].text = ioc.get("source", "N/A")
            for c in cells:
                for paragraph in c.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

        if context["evidence"]["enrichment_summary"]:
            doc.add_paragraph()
            ep = doc.add_paragraph()
            ep.add_run("Enrichment Summary: ").bold = True
            doc.add_paragraph(context["evidence"]["enrichment_summary"])

    # === MITRE ATT&CK MAPPING ===
    doc.add_page_break()
    _add_heading(doc, "MITRE ATT&CK Mapping", level=1)

    if context["mitre"]["kill_chain_phase"]:
        kp = doc.add_paragraph()
        kp.add_run("Kill Chain Phase: ").bold = True
        kp.add_run(context["mitre"]["kill_chain_phase"].replace("_", " ").title())

    if context["mitre"]["techniques"]:
        _add_heading(doc, "Techniques", level=2)
        tt = doc.add_table(rows=1, cols=4)
        tt.style = "Table Grid"
        for i, h in enumerate(["Technique ID", "Name", "Tactic", "Confidence"]):
            tt.rows[0].cells[i].text = h
            for paragraph in tt.rows[0].cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = COLOR_DEEP_NAVY
                    run.font.size = Pt(10)
        for tech in context["mitre"]["techniques"]:
            cells = tt.add_row().cells
            cells[0].text = tech.get("technique_id", "")
            cells[1].text = tech.get("technique_name", "")
            cells[2].text = tech.get("tactic", "")
            cells[3].text = f"{tech.get('confidence', 0) * 100:.0f}%"
            for c in cells:
                for paragraph in c.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

    if context["mitre"]["attack_timeline"]:
        _add_heading(doc, "Attack Timeline", level=2)
        tl = doc.add_table(rows=1, cols=3)
        tl.style = "Table Grid"
        for i, h in enumerate(["Timestamp", "Event", "Technique"]):
            tl.rows[0].cells[i].text = h
            for paragraph in tl.rows[0].cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = COLOR_DEEP_NAVY
                    run.font.size = Pt(10)
        for ev in context["mitre"]["attack_timeline"]:
            cells = tl.add_row().cells
            cells[0].text = ev.get("timestamp", "")
            cells[1].text = ev.get("event", "")
            cells[2].text = ev.get("technique", "")
            for c in cells:
                for paragraph in c.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

    # === DETECTION ENGINEERING ===
    if context["detection"]["sigma_rule"]:
        doc.add_page_break()
        _add_heading(doc, "Detection Engineering", level=1)
        _add_info_table(doc, [
            ("Confidence", f"{context['detection']['confidence'] * 100:.0f}%"),
            ("False Positive Risk", context["detection"]["false_positive_risk"].upper()),
        ])

        if context["detection"]["mitre_techniques_mapped"]:
            doc.add_paragraph()
            tp = doc.add_paragraph()
            tp.add_run("Mapped MITRE Techniques: ").bold = True
            tp.add_run(", ".join(context["detection"]["mitre_techniques_mapped"]))

        if context["detection"]["recommended_log_sources"]:
            doc.add_paragraph()
            lp = doc.add_paragraph()
            lp.add_run("Recommended Log Sources: ").bold = True
            lp.add_run(", ".join(context["detection"]["recommended_log_sources"]))

        _add_heading(doc, "Sigma Rule", level=2)
        sp = doc.add_paragraph()
        sr = sp.add_run(context["detection"]["sigma_rule"])
        sr.font.name = "Courier New"
        sr.font.size = Pt(9)

        if context["detection"]["detection_logic"]:
            _add_heading(doc, "Detection Logic", level=2)
            doc.add_paragraph(context["detection"]["detection_logic"])

    # === RESPONSE PLAYBOOK ===
    doc.add_page_break()
    _add_heading(doc, "Response Playbook", level=1)

    if context["response"]["playbook_name"]:
        _add_info_table(doc, [
            ("Playbook", context["response"]["playbook_name"]),
            ("Priority", context["response"]["priority"].upper()),
            ("Estimated Time", context["response"]["estimated_containment_time"] or "N/A"),
            ("Status", context["response"]["containment_status"].replace("_", " ").title()),
        ])

    if context["response"]["steps"]:
        _add_heading(doc, "Containment Steps", level=2)
        for step in context["response"]["steps"]:
            sp = doc.add_paragraph()
            sp.add_run(f"{step.get('order', 0)}. ").bold = True
            sp.add_run(step.get("action", "")).bold = True

            details = []
            if step.get("tool"):
                details.append(f"Tool: {step['tool']}")
            if step.get("risk_level"):
                details.append(f"Risk: {step['risk_level'].upper()}")
            if step.get("automated"):
                details.append("Automated")
            if details:
                dp = doc.add_paragraph()
                dp.add_run(", ".join(details))
                dp.paragraph_format.left_indent = Inches(0.25)
            if step.get("details"):
                ddp = doc.add_paragraph(step["details"])
                ddp.paragraph_format.left_indent = Inches(0.25)
                ddp.paragraph_format.space_after = Pt(8)

    if context["response"]["post_incident"]:
        _add_heading(doc, "Post-Incident Actions", level=2)
        for action in context["response"]["post_incident"]:
            doc.add_paragraph(action, style="List Bullet")

    # === VALIDATOR AUDIT ===
    if context["validator"]:
        doc.add_page_break()
        _add_heading(doc, "Adversarial Audit", level=1)
        _add_info_table(doc, [
            ("Validation Status", "APPROVED" if context["validator"]["is_approved"] else "REJECTED"),
            ("Risk Score", f"{context['validator']['risk_score'] * 100:.0f}%"),
        ])

        if context["validator"]["critic_comments"]:
            doc.add_paragraph()
            cp = doc.add_paragraph()
            cp.add_run("Critic Comments: ").bold = True
            doc.add_paragraph(context["validator"]["critic_comments"])

        if context["validator"]["safe_alternatives"]:
            _add_heading(doc, "Safer Alternatives", level=2)
            for alt in context["validator"]["safe_alternatives"]:
                doc.add_paragraph(alt, style="List Bullet")

        if context["validator"]["sigma_rule"]:
            _add_heading(doc, "Validator-Generated Sigma Rule", level=2)
            vsp = doc.add_paragraph()
            vsr = vsp.add_run(context["validator"]["sigma_rule"])
            vsr.font.name = "Courier New"
            vsr.font.size = Pt(9)

    # === AUDIT TRAIL ===
    if context["audit_trail"]:
        doc.add_page_break()
        _add_heading(doc, "Audit Trail", level=1)
        at = doc.add_table(rows=1, cols=5)
        at.style = "Table Grid"
        for i, h in enumerate(["Time", "Agent", "Step", "Status", "Time (ms)"]):
            at.rows[0].cells[i].text = h
            for paragraph in at.rows[0].cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = COLOR_DEEP_NAVY
                    run.font.size = Pt(10)
        for entry in context["audit_trail"]:
            ts = entry.get("timestamp", "")
            if "T" in ts:
                ts = ts.split("T")[1].split(".")[0]
            cells = at.add_row().cells
            cells[0].text = ts
            cells[1].text = entry.get("agent", "")
            cells[2].text = entry.get("step", "")
            cells[3].text = entry.get("status", "").upper()
            cells[4].text = str(entry.get("processing_time_ms", 0))
            for c in cells:
                for paragraph in c.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(9)

    # === FOOTER ===
    doc.add_paragraph()
    fp = doc.add_paragraph()
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fr = fp.add_run(
        f"This report was generated automatically by SOCsentinel's Multi-Agent AI System.\n"
        f"Powered by Qwen3 on AMD MI300X (ROCm) | MITRE ATT&CK v16 | "
        f"Investigation ID: {context['investigation_id']}"
    )
    fr.font.size = Pt(9)
    fr.font.color.rgb = COLOR_GRAY
    fr.italic = True

    # Save to buffer
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    logger.debug(
        "DOCX generated successfully",
        investigation_id=context.get("investigation_id"),
        size_bytes=buffer.getbuffer().nbytes,
    )

    return buffer
